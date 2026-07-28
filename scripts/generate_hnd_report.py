"""Generate original ND softcopy Word report for AttendPro (PIN + photo verification)."""
from __future__ import annotations

import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

STUDENT_NAME = "MAKANJUOLA AKINLEKAN OLADELE"
MATRIC_NO = "TEMP/MCET/003"
DEPARTMENT = "COMPUTER SCIENCE"
SCHOOL = "MCET STUDY CENTRE"
POLYTECHNIC = "IGBAJO POLYTECHNIC, IGBAJO, OSUN STATE"
PROGRAMME = "NATIONAL DIPLOMA (ND) IN COMPUTER SCIENCE"
SUPERVISOR = "ALAYANDE, KAYODE GABRIEL"
SESSION = "JULY, 2026"
DEDICATION_FAMILY = "my beloved parents and family"
FRIENDS = "my classmates and friends"

TOPIC = (
    "DESIGN AND IMPLEMENTATION OF A COMPUTERISED STAFF ATTENDANCE "
    "SYSTEM USING PIN AND PHOTO VERIFICATION"
)

OUTPUT = r"C:\Users\Dell\Downloads\MAKANJUOLA_AKINLEKAN_STAFF_ATTENDANCE_PROJECT.docx"
OUTPUT2 = (
    r"C:\Users\Dell\Projects\staff-attendance-system\docs\MAKANJUOLA_AKINLEKAN_STAFF_ATTENDANCE_PROJECT.docx"
)


def set_run_font(run, size=12, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def add_para(
    doc,
    text="",
    *,
    size=12,
    bold=False,
    align="justify",
    space_after=8,
    space_before=0,
    first_line=True,
):
    p = doc.add_paragraph()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "left":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    if first_line and align == "justify" and text:
        pf.first_line_indent = Inches(0.5)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def h_center(doc, text, size=14):
    return add_para(doc, text, size=size, bold=True, align="center", first_line=False, space_after=12)


def h_left(doc, text):
    return add_para(doc, text, bold=True, align="left", first_line=False, space_before=12, space_after=8)


def page_break(doc):
    doc.add_page_break()


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(2.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return doc


def title_page(doc):
    for _ in range(2):
        add_para(doc, "", first_line=False)
    h_center(doc, TOPIC, 14)
    add_para(doc, "", first_line=False)
    add_para(doc, "BY", align="center", bold=True, first_line=False)
    add_para(doc, "", first_line=False)
    h_center(doc, STUDENT_NAME, 13)
    h_center(doc, MATRIC_NO, 12)
    add_para(doc, "", first_line=False)
    add_para(doc, "A PROJECT SUBMITTED TO THE", align="center", bold=True, first_line=False)
    add_para(doc, f"DEPARTMENT OF {DEPARTMENT}", align="center", bold=True, first_line=False)
    add_para(doc, SCHOOL, align="center", bold=True, first_line=False)
    h_center(doc, POLYTECHNIC, 12)
    add_para(doc, "", first_line=False)
    add_para(
        doc,
        "IN PARTIAL FULFILMENT OF THE REQUIREMENTS FOR THE AWARD OF",
        align="center",
        first_line=False,
    )
    h_center(doc, PROGRAMME, 12)
    add_para(doc, "", first_line=False)
    h_center(doc, SESSION, 12)
    page_break(doc)


def dedication(doc):
    h_center(doc, "DEDICATION")
    add_para(
        doc,
        f"This project is dedicated to Almighty God for His guidance throughout my studies, "
        f"and to {DEDICATION_FAMILY}, whose support made this work possible.",
    )
    page_break(doc)


def acknowledgement(doc):
    h_center(doc, "ACKNOWLEDGEMENT")
    add_para(
        doc,
        "I thank Almighty God for the wisdom and strength to complete this project.",
    )
    add_para(
        doc,
        f"I am grateful to my project supervisor, {SUPERVISOR}, for guidance and constructive "
        f"corrections. I also thank the lecturers in the Department of {DEPARTMENT}, "
        f"{POLYTECHNIC}, for the training I received during the National Diploma programme.",
    )
    add_para(
        doc,
        f"My appreciation goes to my family and to {FRIENDS} for encouragement during the "
        f"development and documentation of this system.",
    )
    page_break(doc)


def abstract(doc):
    h_center(doc, "ABSTRACT")
    add_para(
        doc,
        "Staff attendance is a core administrative activity in schools, offices and other "
        "organisations. Many institutions still rely on paper registers. That approach is "
        "slow, difficult to summarise into reports, and vulnerable to buddy punching, where "
        "one person marks attendance for another. This project presents a computerised staff "
        "attendance system that uses Personal Identification Number (PIN) authentication "
        "together with photographic verification at a shared reception kiosk.",
    )
    add_para(
        doc,
        "In the implemented system, a staff member selects his or her name, enters a "
        "four-digit PIN set by the administrator, and captures a live photograph before "
        "clocking in or out. Where a photograph is missing, where a duplicate attempt occurs "
        "on the same day, or where no profile photograph is on file, the attempt is placed "
        "in a review queue for the administrator to compare the live capture with the stored "
        "profile photograph. The application also supports leave requests, attendance "
        "history and management reports.",
    )
    add_para(
        doc,
        "The software was developed with Next.js and TypeScript for the interface, Supabase "
        "for authentication and PostgreSQL storage, and Electron for an optional Windows "
        "desktop package. Testing showed that the system records attendance events when "
        "connected to the internet, blocks self check-in from personal phones, and gives "
        "administrators visual evidence for disputed clocking attempts. The work therefore "
        "offers a practical, low-cost alternative to expensive fingerprint machines while "
        "still strengthening attendance integrity through PIN and photo verification.",
    )
    page_break(doc)


def toc(doc):
    h_center(doc, "TABLE OF CONTENTS")
    for item in [
        "Title Page",
        "Dedication",
        "Acknowledgement",
        "Abstract",
        "Table of Contents",
        "List of Figures",
        "List of Tables",
        "",
        "CHAPTER ONE: INTRODUCTION",
        "1.1 Background of the Study",
        "1.2 Statement of the Problem",
        "1.3 Aim and Objectives of the Study",
        "1.4 Significance of the Study",
        "1.5 Scope of the Study",
        "1.6 Limitations of the Study",
        "1.7 Definition of Terms",
        "",
        "CHAPTER TWO: LITERATURE REVIEW",
        "2.1 Introduction",
        "2.2 Staff Attendance Management",
        "2.3 Manual Attendance Practices",
        "2.4 Related Computerised Approaches",
        "2.5 PIN-Based Authentication",
        "2.6 Photographic Verification in Attendance",
        "2.7 Summary of Reviewed Works",
        "",
        "CHAPTER THREE: SYSTEM ANALYSIS AND METHODOLOGY",
        "3.1 Introduction",
        "3.2 Methodology",
        "3.3 Analysis of the Existing System",
        "3.4 Weaknesses of the Existing System",
        "3.5 The Proposed System",
        "3.6 Requirements Specification",
        "3.7 Feasibility Considerations",
        "",
        "CHAPTER FOUR: SYSTEM DESIGN AND IMPLEMENTATION",
        "4.1 Introduction",
        "4.2 System Architecture",
        "4.3 Database Design",
        "4.4 User Interface Design",
        "4.5 Clocking Process Design",
        "4.6 Implementation Tools",
        "4.7 Implementation Description",
        "4.8 Testing and Results",
        "4.9 Security Measures",
        "",
        "CHAPTER FIVE: SUMMARY, CONCLUSION AND RECOMMENDATIONS",
        "5.1 Summary",
        "5.2 Conclusion",
        "5.3 Recommendations",
        "5.4 Suggestions for Further Studies",
        "",
        "References",
        "Appendix A: Selected Source Code",
        "Appendix B: System Screenshots",
    ]:
        add_para(doc, item, align="left", first_line=False, space_after=2)
    page_break(doc)
    h_center(doc, "LIST OF FIGURES")
    for f in [
        "Figure 3.1 Software Development Stages Used in the Study",
        "Figure 4.1 Overall System Architecture",
        "Figure 4.2 Relationship among Major Data Stores",
        "Figure 4.3 Staff Clocking Flowchart",
        "Figure 4.4 Sign-in Screen",
        "Figure 4.5 Admin Dashboard",
        "Figure 4.6 Staff Record and PIN Setup Screen",
        "Figure 4.7 Reception Kiosk Clocking Screen",
        "Figure 4.8 Attendance Review Queue",
        "Figure 4.9 Attendance Report View",
    ]:
        add_para(doc, f, align="left", first_line=False, space_after=2)
    page_break(doc)
    h_center(doc, "LIST OF TABLES")
    for t in [
        "Table 3.1 Functional Requirements of the Proposed System",
        "Table 3.2 Hardware Requirements",
        "Table 3.3 Software Requirements",
        "Table 4.1 Staff Profile Fields",
        "Table 4.2 Attendance Event Fields",
        "Table 4.3 Review Queue Fields",
        "Table 4.4 Sample Test Results",
    ]:
        add_para(doc, t, align="left", first_line=False, space_after=2)
    page_break(doc)


def chapter_one(doc):
    h_center(doc, "CHAPTER ONE")
    h_center(doc, "INTRODUCTION")
    add_para(
        doc,
        "This chapter introduces the project. It explains why a computerised attendance "
        "system with PIN and photo verification is needed, states the research problem, "
        "lists the objectives, and defines the scope and key terms used in the report.",
    )

    h_left(doc, "1.1 BACKGROUND OF THE STUDY")
    add_para(
        doc,
        "Every organisation that pays wages or monitors duty hours needs a trustworthy "
        "way of knowing when staff arrive and when they leave. Attendance data supports "
        "payroll, discipline and everyday planning. Where records are weak, lateness and "
        "absenteeism become hard to control, and disputes are common.",
    )
    add_para(
        doc,
        "Paper attendance books remain common in many Nigerian workplaces because they "
        "are familiar and cheap. Unfortunately, they also allow buddy punching, create "
        "unclear handwriting, and force clerks to spend long hours preparing monthly "
        "summaries. Computerised systems were introduced to reduce these problems. Some "
        "organisations buy fingerprint machines; others use smart cards or mobile apps. "
        "Fingerprint devices work well but are costly for small centres. Mobile self "
        "check-in is convenient but easy to abuse when staff clock in from home.",
    )
    add_para(
        doc,
        "This project therefore focuses on a middle path: a shared reception kiosk where "
        "each staff member proves identity with a secret four-digit PIN and leaves a "
        "photograph that management can inspect. The resulting system, AttendPro, stores "
        "records online so administrators can view live attendance, manage leave and "
        "export reports without hunting through paper files.",
    )

    h_left(doc, "1.2 STATEMENT OF THE PROBLEM")
    add_para(doc, "The problems that motivated this project include the following:", first_line=False)
    for item in [
        "Buddy punching on paper registers is easy and hard to prove later.",
        "Manual summaries for the month are slow and often contain arithmetic mistakes.",
        "Lost or damaged attendance books destroy historical evidence.",
        "Fingerprint scanners are effective but too expensive for many small institutions.",
        "Allowing staff to clock in from personal phones weakens location and identity control.",
        "Administrators need visual evidence when a disputed clocking event is reported.",
    ]:
        add_para(doc, item, align="left", first_line=False)

    h_left(doc, "1.3 AIM AND OBJECTIVES OF THE STUDY")
    add_para(
        doc,
        "The aim of the study is to design and implement a computerised staff attendance "
        "system that authenticates workers with a PIN and verifies each clocking event "
        "with a photograph.",
    )
    add_para(doc, "The objectives are to:", first_line=False)
    for i, o in enumerate(
        [
            "Provide a shared kiosk for clock-in and clock-out.",
            "Authenticate each staff member with an administrator-issued four-digit PIN.",
            "Capture a photograph at every successful clocking attempt.",
            "Route doubtful attempts to an administrative review queue.",
            "Support staff and admin portals for leave, history and reports.",
            "Package the solution for web use and optional Windows desktop installation.",
        ],
        1,
    ):
        add_para(doc, f"{i}. {o}", align="left", first_line=False)

    h_left(doc, "1.4 SIGNIFICANCE OF THE STUDY")
    add_para(
        doc,
        "The study is useful to school and office administrators who need stronger "
        "attendance control without buying biometric hardware. It is useful to staff "
        "because they can view their own records online. It is useful to students of "
        "computer science as a complete example of analysis, design, implementation and "
        "testing of a multi-user web information system. Finally, it reduces dependence "
        "on fragile paper records.",
    )

    h_left(doc, "1.5 SCOPE OF THE STUDY")
    add_para(
        doc,
        "The project covers staff registration, PIN assignment, profile photograph "
        "upload, kiosk clocking with live photo capture, review of exceptional attempts, "
        "leave request handling, attendance reporting, and deployment as a web or "
        "desktop application for an organisation such as a study centre or office.",
    )

    h_left(doc, "1.6 LIMITATIONS OF THE STUDY")
    for lim in [
        "An internet connection is required for authentication and data storage.",
        "Comparison of photographs is done by the administrator, not by automatic face matching in the current release.",
        "Staff cannot clock in from personal phones; only the configured kiosk is used.",
        "The project period limited how many organisations could be used for field trials.",
    ]:
        add_para(doc, lim, align="left", first_line=False)

    h_left(doc, "1.7 DEFINITION OF TERMS")
    terms = {
        "Attendance": "Record of a staff member's presence at work during duty hours.",
        "Clock-in / Clock-out": "Electronic recording of arrival time / departure time.",
        "PIN": "A four-digit secret code used to confirm staff identity at the kiosk.",
        "Photo verification": "Use of a captured image as evidence that the correct person clocked.",
        "Kiosk": "A shared computer or tablet kept at reception for attendance.",
        "Buddy punching": "Marking attendance for another person who is not present.",
        "Review queue": "List of clocking attempts waiting for administrator decision.",
        "Database": "Organised electronic store of staff and attendance data.",
    }
    for k, v in terms.items():
        add_para(doc, f"{k}: {v}", align="left", first_line=False)
    page_break(doc)


def chapter_two(doc):
    h_center(doc, "CHAPTER TWO")
    h_center(doc, "LITERATURE REVIEW")
    h_left(doc, "2.1 INTRODUCTION")
    add_para(
        doc,
        "This chapter reviews ideas and systems related to staff attendance. The goal "
        "is not to copy an older fingerprint project, but to place PIN and photographic "
        "verification within the wider body of attendance technology and to identify "
        "gaps that the present work addresses.",
    )
    h_left(doc, "2.2 STAFF ATTENDANCE MANAGEMENT")
    add_para(
        doc,
        "Attendance management is part of human resource administration. Accurate "
        "presence data helps organisations calculate hours worked, enforce lateness "
        "rules and plan duty rosters. Scholars of management information systems note "
        "that timely and reliable operational data improves decision making. Attendance "
        "is therefore both an administrative routine and an information-system problem.",
    )
    h_left(doc, "2.3 MANUAL ATTENDANCE PRACTICES")
    add_para(
        doc,
        "Manual registers require each worker to write a name, time and signature. The "
        "method needs no electricity and little training, yet it scales poorly. Large "
        "staff lists make checking tedious. Paper is also easy to alter and difficult "
        "to back up. These weaknesses justify computerisation even when budgets are "
        "small.",
    )
    h_left(doc, "2.4 RELATED COMPUTERISED APPROACHES")
    add_para(
        doc,
        "Published and commercial systems fall into several families: barcode or RFID "
        "card punch systems; fingerprint and face biometric clocks; and browser or "
        "mobile applications. Card systems are fast but cards can be shared. Biometric "
        "clocks raise security but raise cost. Pure mobile apps raise convenience but "
        "weaken control of place and identity unless extra checks are added.",
    )
    h_left(doc, "2.5 PIN-BASED AUTHENTICATION")
    add_para(
        doc,
        "A Personal Identification Number is a short shared secret. Banks and access-"
        "control systems have long used PINs because they are simple to remember and "
        "cheap to implement. In attendance, a PIN alone is not enough if staff can "
        "whisper it to a friend. Combining a PIN with a controlled physical location "
        "(the kiosk) and a photograph reduces that risk without requiring fingerprint "
        "sensors.",
    )
    h_left(doc, "2.6 PHOTOGRAPHIC VERIFICATION IN ATTENDANCE")
    add_para(
        doc,
        "A photograph taken at the moment of clocking creates an audit trail. Even when "
        "automatic face matching is not used, an administrator can later compare the "
        "live image with a stored profile photograph. This manual review model is "
        "suitable for student projects and small organisations because it avoids "
        "expensive specialised hardware while still providing evidence.",
    )
    h_left(doc, "2.7 SUMMARY OF REVIEWED WORKS")
    add_para(
        doc,
        "The review shows that fingerprint attendance is strong but costly, while "
        "uncontrolled mobile check-in is weak against fraud. A kiosk that demands a "
        "PIN and a photograph sits between those extremes. That observation guided the "
        "design choices documented in Chapters Three and Four.",
    )
    page_break(doc)


def chapter_three(doc):
    h_center(doc, "CHAPTER THREE")
    h_center(doc, "SYSTEM ANALYSIS AND METHODOLOGY")
    h_left(doc, "3.1 INTRODUCTION")
    add_para(
        doc,
        "This chapter explains how the existing attendance practice was studied, which "
        "development method was used, and what the proposed PIN-and-photo system must "
        "do.",
    )
    h_left(doc, "3.2 METHODOLOGY")
    add_para(
        doc,
        "A staged software process similar to the classical waterfall approach was "
        "followed: clarify requirements, analyse the current practice, design the "
        "solution, implement modules, test them, and document the result. Observation "
        "of paper-based attendance and study of related systems provided the facts used "
        "in analysis.",
    )
    add_para(
        doc,
        "[INSERT FIGURE 3.1 — Software Development Stages Used in the Study]",
        align="center",
        first_line=False,
        bold=True,
    )

    h_left(doc, "3.3 ANALYSIS OF THE EXISTING SYSTEM")
    add_para(
        doc,
        "In the existing arrangement, a register is kept at the entrance. Staff write "
        "arrival and departure times by hand. At the end of the month an officer counts "
        "entries and prepares a summary for management. Leave notes are often filed "
        "separately, so combining leave and attendance is awkward.",
    )

    h_left(doc, "3.4 WEAKNESSES OF THE EXISTING SYSTEM")
    for w in [
        "Identity of the person who wrote an entry is hard to confirm.",
        "Monthly compilation consumes staff time.",
        "Paper records are fragile.",
        "Managers cannot see who is present in real time.",
        "There is no automatic alert for unusual repeated entries.",
    ]:
        add_para(doc, w, align="left", first_line=False)

    h_left(doc, "3.5 THE PROPOSED SYSTEM")
    add_para(
        doc,
        "AttendPro replaces the register with an electronic workflow. The administrator "
        "creates staff accounts, sets each person's PIN and uploads a profile photograph. "
        "At the kiosk, the staff member picks a name, types the PIN and takes a photo. "
        "Valid events become attendance records. Doubtful events wait in a review queue. "
        "Staff may view history and request leave through a portal; administrators manage "
        "people, leave and reports.",
    )

    h_left(doc, "3.6 REQUIREMENTS SPECIFICATION")
    add_para(doc, "Table 3.1 — Functional Requirements of the Proposed System", bold=True, align="left", first_line=False)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "S/N"
    table.rows[0].cells[1].text = "Requirement"
    for i, r in enumerate(
        [
            "Register and maintain staff profiles.",
            "Assign and update a four-digit kiosk PIN per staff member.",
            "Store a profile photograph for visual comparison.",
            "Allow clock-in/out only through an authenticated kiosk session.",
            "Require PIN and photo for each clocking attempt.",
            "Queue missing-photo, duplicate-day and no-profile-photo cases for review.",
            "Allow administrators to approve or reject queued items.",
            "Support leave requests and attendance reporting.",
        ],
        1,
    ):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = r
    add_para(doc, "", first_line=False)

    add_para(doc, "Table 3.2 — Hardware Requirements", bold=True, align="left", first_line=False)
    hw = doc.add_table(rows=1, cols=3)
    hw.style = "Table Grid"
    hw.rows[0].cells[0].text = "S/N"
    hw.rows[0].cells[1].text = "Item"
    hw.rows[0].cells[2].text = "Minimum"
    for i, (a, b) in enumerate(
        [
            ("Computer / tablet", "Dual-core CPU, 4 GB RAM"),
            ("Camera", "Webcam for kiosk photo capture"),
            ("Network", "Internet access"),
            ("Storage", "Enough space for local desktop install"),
        ],
        1,
    ):
        row = hw.add_row().cells
        row[0].text = str(i)
        row[1].text = a
        row[2].text = b
    add_para(doc, "", first_line=False)

    add_para(doc, "Table 3.3 — Software Requirements", bold=True, align="left", first_line=False)
    sw = doc.add_table(rows=1, cols=3)
    sw.style = "Table Grid"
    sw.rows[0].cells[0].text = "S/N"
    sw.rows[0].cells[1].text = "Software"
    sw.rows[0].cells[2].text = "Role"
    for i, (a, b) in enumerate(
        [
            ("Windows 10/11 or modern OS", "Host environment"),
            ("Web browser", "Access web application"),
            ("AttendPro desktop build", "Optional installed client"),
            ("Supabase / PostgreSQL", "Auth, database, file storage"),
        ],
        1,
    ):
        row = sw.add_row().cells
        row[0].text = str(i)
        row[1].text = a
        row[2].text = b

    h_left(doc, "3.7 FEASIBILITY CONSIDERATIONS")
    add_para(
        doc,
        "Technically, the tools used are mainstream and well documented. Economically, "
        "the design avoids fingerprint readers. Operationally, the kiosk steps are short "
        "enough for everyday use after brief training. The proposal is therefore "
        "feasible for an ND project and for small institutions.",
    )
    page_break(doc)


def chapter_four(doc):
    h_center(doc, "CHAPTER FOUR")
    h_center(doc, "SYSTEM DESIGN AND IMPLEMENTATION")
    h_left(doc, "4.1 INTRODUCTION")
    add_para(
        doc,
        "This chapter describes how the PIN-and-photo attendance system was designed "
        "and built, how data are organised, how clocking works, and how the software "
        "was tested.",
    )

    h_left(doc, "4.2 SYSTEM ARCHITECTURE")
    add_para(
        doc,
        "Clients (browser or Electron window) talk to application routes and server "
        "actions. Persistent data live in Supabase PostgreSQL. Photographs are kept in "
        "private storage buckets. Kiosk devices authenticate with an API key and hold "
        "a session cookie. Staff accounts use email and password for the portal, while "
        "kiosk clocking uses the PIN rather than the portal password.",
    )
    add_para(doc, "[INSERT FIGURE 4.1 — Overall System Architecture]", align="center", first_line=False, bold=True)

    h_left(doc, "4.3 DATABASE DESIGN")
    add_para(
        doc,
        "Major tables include organisation profiles, staff profiles (with hashed PIN "
        "and avatar path), kiosk devices, attendance event records, review-queue items "
        "and leave requests. Row Level Security keeps each organisation's data separate.",
    )
    add_para(doc, "Table 4.1 — Staff Profile Fields", bold=True, align="left", first_line=False)
    t1 = doc.add_table(rows=1, cols=3)
    t1.style = "Table Grid"
    t1.rows[0].cells[0].text = "Field"
    t1.rows[0].cells[1].text = "Type"
    t1.rows[0].cells[2].text = "Purpose"
    for a, b, c in [
        ("full_name", "Text", "Display name on kiosk list"),
        ("email", "Text", "Portal login"),
        ("kiosk_pin_hash", "Text", "Hashed PIN for kiosk"),
        ("avatar_url", "Text", "Stored profile photograph"),
        ("role", "Text", "admin or staff"),
        ("is_active", "Boolean", "Block inactive accounts"),
    ]:
        row = t1.add_row().cells
        row[0].text = a
        row[1].text = b
        row[2].text = c
    add_para(doc, "", first_line=False)
    add_para(doc, "Table 4.2 — Attendance Event Fields", bold=True, align="left", first_line=False)
    t2 = doc.add_table(rows=1, cols=3)
    t2.style = "Table Grid"
    t2.rows[0].cells[0].text = "Field"
    t2.rows[0].cells[1].text = "Type"
    t2.rows[0].cells[2].text = "Purpose"
    for a, b, c in [
        ("type", "Text", "check_in or check_out"),
        ("server_timestamp", "Timestamp", "Official event time"),
        ("photo_capture_url", "Text", "Live kiosk photograph"),
        ("match_status", "Text", "auto or manual approval"),
        ("kiosk_device_id", "UUID", "Which kiosk was used"),
    ]:
        row = t2.add_row().cells
        row[0].text = a
        row[1].text = b
        row[2].text = c
    add_para(doc, "", first_line=False)
    add_para(doc, "[INSERT FIGURE 4.2 — Relationship among Major Data Stores]", align="center", first_line=False, bold=True)

    h_left(doc, "4.4 USER INTERFACE DESIGN")
    add_para(
        doc,
        "Administrators see a dashboard, staff list, leave page, review queue, reports "
        "and settings. Staff see attendance history, leave forms and profile photo "
        "upload. The kiosk interface is deliberately simple: search or pick a name, "
        "enter PIN, take photo, then view the result message.",
    )

    h_left(doc, "4.5 CLOCKING PROCESS DESIGN")
    add_para(
        doc,
        "After the staff member is selected, the PIN is verified against the stored "
        "hash. A photograph must be uploaded. The server then checks for cooldown, "
        "same-day duplicates and missing profile photographs. Clean attempts create an "
        "attendance record immediately. Other attempts enter the review queue with both "
        "images available to the administrator.",
    )
    add_para(doc, "[INSERT FIGURE 4.3 — Staff Clocking Flowchart]", align="center", first_line=False, bold=True)

    h_left(doc, "4.6 IMPLEMENTATION TOOLS")
    add_para(
        doc,
        "Next.js and TypeScript powered the application screens and server logic. "
        "Tailwind CSS styled the interface. Supabase provided authentication, database "
        "and storage. Electron packaged a Windows installer for softcopy demonstration. "
        "Git was used for version control.",
    )

    h_left(doc, "4.7 IMPLEMENTATION DESCRIPTION")
    add_para(
        doc,
        "Modules were coded for organisation and staff management, kiosk session "
        "creation, photo upload, PIN verification, clock processing, review resolution, "
        "leave workflow and reporting. Demo credentials and a hosted web deployment "
        "support viva demonstration alongside the desktop installer.",
    )
    for fig in [
        "[INSERT FIGURE 4.4 — Sign-in Screen]",
        "[INSERT FIGURE 4.5 — Admin Dashboard]",
        "[INSERT FIGURE 4.6 — Staff Record and PIN Setup Screen]",
        "[INSERT FIGURE 4.7 — Reception Kiosk Clocking Screen]",
        "[INSERT FIGURE 4.8 — Attendance Review Queue]",
        "[INSERT FIGURE 4.9 — Attendance Report View]",
    ]:
        add_para(doc, fig, align="center", first_line=False, bold=True)

    h_left(doc, "4.8 TESTING AND RESULTS")
    add_para(doc, "Table 4.4 — Sample Test Results", bold=True, align="left", first_line=False)
    tt = doc.add_table(rows=1, cols=4)
    tt.style = "Table Grid"
    tt.rows[0].cells[0].text = "S/N"
    tt.rows[0].cells[1].text = "Scenario"
    tt.rows[0].cells[2].text = "Expected outcome"
    tt.rows[0].cells[3].text = "Result"
    for i, (a, b, c) in enumerate(
        [
            ("Correct portal login", "Dashboard or staff home opens", "Passed"),
            ("Wrong portal password", "Access denied", "Passed"),
            ("Correct PIN + photo", "Clock event saved", "Passed"),
            ("Wrong PIN", "Rejection message", "Passed"),
            ("No photo", "Item sent to review", "Passed"),
            ("Duplicate same-day action", "Item sent to review", "Passed"),
            ("Approve review item", "Attendance recorded", "Passed"),
            ("Leave request by staff", "Pending leave created", "Passed"),
        ],
        1,
    ):
        row = tt.add_row().cells
        row[0].text = str(i)
        row[1].text = a
        row[2].text = b
        row[3].text = c

    h_left(doc, "4.9 SECURITY MEASURES")
    for s in [
        "PINs are hashed before storage.",
        "Organisation data are isolated with Row Level Security.",
        "Only administrators manage staff, PINs and reviews.",
        "Self check-in from personal devices is disabled; kiosk clocking is required.",
        "Photographs are stored privately and shown through controlled links.",
        "Transport uses HTTPS.",
    ]:
        add_para(doc, s, align="left", first_line=False)
    page_break(doc)


def chapter_five(doc):
    h_center(doc, "CHAPTER FIVE")
    h_center(doc, "SUMMARY, CONCLUSION AND RECOMMENDATIONS")
    h_left(doc, "5.1 SUMMARY")
    add_para(
        doc,
        "A computerised staff attendance system based on PIN and photo verification was "
        "analysed, designed and implemented. The kiosk workflow, review queue, leave "
        "module and reporting features were tested successfully under normal network "
        "conditions.",
    )
    h_left(doc, "5.2 CONCLUSION")
    add_para(
        doc,
        "Paper attendance is no longer adequate where management needs speed and "
        "accountability. This project shows that a practical computerised alternative "
        "can be built without fingerprint hardware by combining a secret PIN, a "
        "controlled kiosk location and photographic evidence. The objectives of the "
        "study were met.",
    )
    h_left(doc, "5.3 RECOMMENDATIONS")
    for r in [
        "Institutions should migrate from paper registers to computerised attendance.",
        "Administrators should protect staff PINs and reset them when compromised.",
        "Every staff member should keep an up-to-date profile photograph.",
        "The reception area should have reliable internet for the kiosk.",
    ]:
        add_para(doc, r, align="left", first_line=False)
    h_left(doc, "5.4 SUGGESTIONS FOR FURTHER STUDIES")
    for s in [
        "Integrate automatic face matching (for example a third-party identity API) on every clock-in and clock-out.",
        "Improve offline capture with later synchronisation.",
        "Link attendance totals to payroll calculation.",
        "Extend alerts for lateness through SMS or email.",
    ]:
        add_para(doc, s, align="left", first_line=False)
    page_break(doc)


def references(doc):
    h_center(doc, "REFERENCES")
    for r in [
        "Date, C. J. (2004). An Introduction to Database Systems. Pearson Education.",
        "Laudon, K. C., & Laudon, J. P. (2020). Management Information Systems. Pearson.",
        "Pressman, R. S. (2014). Software Engineering: A Practitioner's Approach. McGraw-Hill.",
        "Sommerville, I. (2016). Software Engineering. Pearson Education.",
        "Stair, R., & Reynolds, G. (2018). Fundamentals of Information Systems. Cengage Learning.",
        "Next.js Documentation. (2024). Retrieved from https://nextjs.org/docs",
        "Supabase Documentation. (2024). Retrieved from https://supabase.com/docs",
        "Electron Documentation. (2024). Retrieved from https://www.electronjs.org/docs",
    ]:
        add_para(doc, r, align="left", first_line=False, space_after=6)
    page_break(doc)
    h_center(doc, "APPENDIX A")
    h_center(doc, "SELECTED SOURCE CODE")
    add_para(
        doc,
        "Attach selected listings for PIN hashing/verification and kiosk clock "
        "processing. Do not include secret keys.",
        first_line=False,
    )
    page_break(doc)
    h_center(doc, "APPENDIX B")
    h_center(doc, "SYSTEM SCREENSHOTS")
    add_para(
        doc,
        "Insert screenshots of login, dashboard, staff PIN setup, kiosk clocking, "
        "review queue and reports.",
        first_line=False,
    )


def main():
    doc = setup_doc()
    title_page(doc)
    dedication(doc)
    acknowledgement(doc)
    abstract(doc)
    toc(doc)
    chapter_one(doc)
    chapter_two(doc)
    chapter_three(doc)
    chapter_four(doc)
    chapter_five(doc)
    references(doc)
    os.makedirs(os.path.dirname(OUTPUT2), exist_ok=True)
    try:
        doc.save(OUTPUT)
        print("Saved:", OUTPUT)
    except PermissionError:
        alt = OUTPUT.replace(".docx", "_v2.docx")
        doc.save(alt)
        print("Downloads file locked; saved:", alt)
    doc.save(OUTPUT2)
    print("Saved:", OUTPUT2)


if __name__ == "__main__":
    main()
